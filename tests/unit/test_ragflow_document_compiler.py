from __future__ import annotations

import ast
import base64
import io
import json
import subprocess
import sys
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path
from types import SimpleNamespace
from typing import Any, cast

import pytest
import rfc8785
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from pypdf import PdfWriter
from pypdf.generic import Destination

from adapters.parsers import ragflow_documents as ragflow_document_adapter
from adapters.parsers.ragflow_documents import compile_document_bytes
from applications.document_compiler_runner import (
    ArtifactSource,
    BytesArtifactSource,
    _document_runner_environment,
    compile_in_local_document_runner,
)
from engine.supply import (
    DOCX_CONFIG_V1,
    MAX_FORMAT_DOCUMENT_UNITS,
    PDF_TEXT_OUTLINE_V1,
    CompilationProfileRef,
    DocumentCompilationFailure,
    DocumentCompilationFailureCode,
    DocumentStructuralKind,
    DocxXmlLocator,
    ParsedDocument,
    PdfRegionLocator,
    StructuralUnit,
    canonicalize_parsed_document,
    deserialize_parsed_document,
)
from eval._compiler_acceptance import acceptance_context
from third_party.ragflow.deepdoc.parser import utils as ragflow_pdf_utils
from third_party.ragflow.deepdoc.parser.utils import RawPdfOutline

REPOSITORY_ROOT = Path(__file__).parents[2]


@dataclass
class _CountingArtifact(ArtifactSource):
    payload: bytes
    reads: int = 0

    def read(self) -> bytes:
        self.reads += 1
        return self.payload


def _docx_fixture(*, with_image: bool = False) -> bytes:
    document = Document()
    document.add_heading("Architecture", level=1)
    document.add_paragraph("First paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "parser"
    table.cell(1, 1).text = "registered"
    document.add_paragraph("Last paragraph.")
    if with_image:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        run = document.add_paragraph().add_run()
        run._r.append(parse_xml(f"<pic:pic {nsdecls('pic')}></pic:pic>"))
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _save_docx(document: DocumentType) -> bytes:
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _docx_with_unsupported_body_container() -> bytes:
    document = Document()
    document.add_paragraph("Retained body text.")
    content_control = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Content control text must not disappear."
    run.append(text)
    paragraph.append(run)
    content.append(paragraph)
    content_control.append(content)
    document.element.body.insert(-1, content_control)
    return _save_docx(document)


def _docx_with_tracked_insertion() -> bytes:
    document = Document()
    document.add_paragraph("Retained body text.")
    paragraph = document.add_paragraph()
    insertion = OxmlElement("w:ins")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Tracked insertion must not disappear."
    run.append(text)
    insertion.append(run)
    paragraph._p.append(insertion)
    return _save_docx(document)


def _docx_with_unsupported_drawing(*, in_header: bool) -> bytes:
    document = Document()
    document.add_paragraph("Retained body text.")
    paragraph = (
        document.sections[0].header.paragraphs[0]
        if in_header
        else document.add_paragraph()
    )
    paragraph.add_run()._r.append(OxmlElement("w:drawing"))
    return _save_docx(document)


def _docx_with_nested_table() -> bytes:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Outer cell"
    nested = table.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Nested cell must not disappear."
    return _save_docx(document)


def _docx_fixture_with_blank_source_block() -> bytes:
    document = Document()
    document.add_heading("Architecture", level=1)
    document.add_paragraph("")
    document.add_paragraph("After blank source block.")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_outline_fixture() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    root = writer.add_outline_item("Overview", 0)
    writer.add_outline_item("Details", 1, parent=root)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def _pdf_outline_fixture_for_same_page(*, shifted: bool = False) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    if shifted:
        page.mediabox.lower_left = (-10, -10)
        page.mediabox.upper_right = (602, 782)
    writer.add_outline_item("First", 0)
    writer.add_outline_item("Second", 0)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def test_unknown_profile_refuses_before_artifact_bytes_are_opened() -> None:
    artifact = _CountingArtifact(b"must not be read")

    outcome = compile_in_local_document_runner(
        artifact,
        "pdf-layout-ocr-v1",
        acceptance_context=acceptance_context(),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.UNKNOWN_PROFILE
    assert artifact.reads == 0


def test_child_unknown_profile_returns_closed_failure() -> None:
    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "applications.document_compiler_runner",
            "--profile",
            "pdf-layout-ocr-v1",
        ],
        input=b"must not be parsed",
        capture_output=True,
        check=True,
        timeout=30,
    )

    assert json.loads(completed.stdout) == {
        "outcome": "failure",
        "failure": {"code": "unknown_profile"},
    }


def test_child_enforces_artifact_bound_with_a_closed_refusal() -> None:
    from applications.document_compiler_runner import MAX_DOCUMENT_ARTIFACT_BYTES

    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "applications.document_compiler_runner",
            "--profile",
            DOCX_CONFIG_V1,
        ],
        input=b"x" * (MAX_DOCUMENT_ARTIFACT_BYTES + 1),
        capture_output=True,
        check=True,
        timeout=30,
    )

    assert json.loads(completed.stdout) == {
        "outcome": "failure",
        "failure": {"code": "artifact_bound_exceeded"},
    }


def test_owned_document_runner_has_no_network_database_or_model_imports() -> None:
    paths = (
        REPOSITORY_ROOT / "adapters/parsers/ragflow_documents.py",
        REPOSITORY_ROOT / "applications/document_compiler_runner.py",
        REPOSITORY_ROOT / "third_party/ragflow/deepdoc/parser/docx_parser.py",
        REPOSITORY_ROOT / "third_party/ragflow/deepdoc/parser/utils.py",
    )
    forbidden = {
        "common",
        "huggingface_hub",
        "httpx",
        "os",
        "psycopg",
        "requests",
        "socket",
        "sqlalchemy",
        "urllib",
    }
    for path in paths:
        imports: set[str] = set()
        tree = ast.parse(path.read_bytes(), filename=str(path))
        for node in ast.walk(tree):
            if isinstance(node, ast.Import):
                imports.update(alias.name.partition(".")[0] for alias in node.names)
            elif isinstance(node, ast.ImportFrom) and node.level == 0:
                assert node.module is not None
                imports.add(node.module.partition(".")[0])
        assert imports.isdisjoint(forbidden), path


def test_docx_profile_preserves_ooxml_block_order_and_typed_locators() -> None:
    source = _docx_fixture()

    outcome = compile_document_bytes(
        source,
        CompilationProfileRef("context-engine-docx-v1", DOCX_CONFIG_V1),
    )

    assert type(outcome) is ParsedDocument
    assert outcome.units is not None
    assert [unit.kind for unit in outcome.units] == [
        DocumentStructuralKind.HEADING,
        DocumentStructuralKind.PARAGRAPH,
        DocumentStructuralKind.TABLE,
        DocumentStructuralKind.PARAGRAPH,
    ]
    assert [unit.text for unit in outcome.units] == [
        "Architecture",
        "First paragraph.",
        "Key\tValue\nparser\tregistered",
        "Last paragraph.",
    ]
    assert all(
        type(locator) is DocxXmlLocator
        for unit in outcome.units
        for locator in unit.locators
    )
    docx_locators = tuple(unit.locators[0] for unit in outcome.units)
    assert all(type(locator) is DocxXmlLocator for locator in docx_locators)
    assert tuple(
        locator.block_ordinal
        for locator in docx_locators
        if type(locator) is DocxXmlLocator
    ) == (0, 1, 2, 3)
    assert outcome.provenance.config_version == DOCX_CONFIG_V1


def test_docx_image_is_an_honest_typed_refusal() -> None:
    outcome = compile_document_bytes(
        _docx_fixture(with_image=True),
        CompilationProfileRef("context-engine-docx-v1", DOCX_CONFIG_V1),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.FIGURE_NOT_SUPPORTED


@pytest.mark.parametrize(
    "source_builder",
    (
        _docx_with_unsupported_body_container,
        _docx_with_tracked_insertion,
        _docx_with_nested_table,
    ),
    ids=("content-control", "tracked-insertion", "nested-table"),
)
def test_docx_refuses_source_content_it_cannot_preserve(
    source_builder: Callable[[], bytes],
) -> None:
    outcome = compile_document_bytes(
        source_builder(),
        CompilationProfileRef("context-engine-docx-v1", DOCX_CONFIG_V1),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.INVALID_ARTIFACT


@pytest.mark.parametrize("in_header", (False, True))
def test_docx_refuses_unsupported_drawings_in_every_package_part(
    in_header: bool,
) -> None:
    outcome = compile_document_bytes(
        _docx_with_unsupported_drawing(in_header=in_header),
        CompilationProfileRef("context-engine-docx-v1", DOCX_CONFIG_V1),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.FIGURE_NOT_SUPPORTED


def test_docx_locator_ordinal_is_the_ooxml_source_ordinal_not_output_index() -> None:
    outcome = compile_document_bytes(
        _docx_fixture_with_blank_source_block(),
        CompilationProfileRef("context-engine-docx-v1", DOCX_CONFIG_V1),
    )

    assert type(outcome) is ParsedDocument
    assert outcome.units is not None
    locators = tuple(unit.locators[0] for unit in outcome.units)
    assert tuple(
        locator.block_ordinal
        for locator in locators
        if type(locator) is DocxXmlLocator
    ) == (0, 2)


def test_pdf_outline_profile_emits_source_order_and_pdf_region_locators() -> None:
    outcome = compile_document_bytes(
        _pdf_outline_fixture(),
        CompilationProfileRef(
            "context-engine-pdf-outline-v1",
            PDF_TEXT_OUTLINE_V1,
        ),
    )

    assert type(outcome) is ParsedDocument
    assert outcome.units is not None
    assert [unit.text for unit in outcome.units] == ["Overview", "Details"]
    assert [unit.heading_level for unit in outcome.units] == [1, 2]
    assert all(unit.kind is DocumentStructuralKind.HEADING for unit in outcome.units)
    assert all(
        type(unit.locators[0]) is PdfRegionLocator for unit in outcome.units
    )
    pdf_locators = tuple(unit.locators[0] for unit in outcome.units)
    assert [
        locator.page_number
        for locator in pdf_locators
        if type(locator) is PdfRegionLocator
    ] == [1, 2]


def test_pdf_outline_normalizes_shifted_media_box_coordinates() -> None:
    outcome = compile_document_bytes(
        _pdf_outline_fixture_for_same_page(shifted=True),
        CompilationProfileRef(
            "context-engine-pdf-outline-v1",
            PDF_TEXT_OUTLINE_V1,
        ),
    )

    assert type(outcome) is ParsedDocument
    assert outcome.units is not None
    locator = outcome.units[0].locators[0]
    assert type(locator) is PdfRegionLocator
    assert locator.bbox_points == (0.0, 0.0, 612.0, 792.0)


def test_pdf_outline_bounds_page_before_rendering(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=20_001, height=20_001)
    writer.add_outline_item("Oversized", 0)
    output = io.BytesIO()
    writer.write(output)

    def reject_render(_page: object) -> str:
        raise AssertionError("oversized PDF page was rendered")

    monkeypatch.setattr(ragflow_pdf_utils, "_page_render_digest", reject_render)
    outcome = compile_document_bytes(
        output.getvalue(),
        CompilationProfileRef(
            "context-engine-pdf-outline-v1",
            PDF_TEXT_OUTLINE_V1,
        ),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.DOCUMENT_BOUND_EXCEEDED


def test_pdf_outline_rejects_non_finite_page_before_rendering(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = _pdf_outline_fixture_for_same_page()
    actual_reader = cast(Any, ragflow_pdf_utils).PdfReader(io.BytesIO(source))

    class NonFinitePage:
        mediabox = (float("nan"), 0.0, 612.0, 792.0)

    class NonFiniteReader:
        outline = actual_reader.outline
        pages = (NonFinitePage(),)

        def get_destination_page_number(self, node: Destination) -> int:
            page_number = actual_reader.get_destination_page_number(node)
            assert type(page_number) is int
            return page_number

    def reject_render(_page: object) -> str:
        raise AssertionError("non-finite PDF page was rendered")

    monkeypatch.setattr(
        ragflow_pdf_utils, "PdfReader", lambda _source: NonFiniteReader()
    )
    monkeypatch.setattr(ragflow_pdf_utils, "_page_render_digest", reject_render)
    outcome = compile_document_bytes(
        source,
        CompilationProfileRef(
            "context-engine-pdf-outline-v1",
            PDF_TEXT_OUTLINE_V1,
        ),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.DOCUMENT_BOUND_EXCEEDED


def test_pdf_locator_constructor_enforces_pixel_area_bound() -> None:
    digest = "0" * 64

    with pytest.raises(ValueError, match="pixel-area hard bound"):
        PdfRegionLocator(
            artifact_digest=digest,
            page_number=1,
            bbox_points=(0.0, 0.0, 10_000.0, 5_000.0),
            page_render_digest=digest,
            extraction_method="pypdf-outline-v1",
        )


def test_pdf_outline_closes_document_and_renders_each_page_once(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    pdfium_module = cast(Any, ragflow_pdf_utils).pdfium
    original_document = pdfium_module.PdfDocument
    render_calls = 0

    class TrackingDocument:
        def __init__(self, source: bytes) -> None:
            self._document = original_document(source)
            self.closed = False
            wrappers.append(self)

        def __getitem__(self, index: int) -> object:
            return self._document[index]

        def close(self) -> None:
            self.closed = True
            self._document.close()

    wrappers: list[TrackingDocument] = []

    def recording_digest(page: object) -> str:
        nonlocal render_calls
        render_calls += 1
        cast(Any, page).close()
        return "0" * 64

    monkeypatch.setattr(pdfium_module, "PdfDocument", TrackingDocument)
    monkeypatch.setattr(ragflow_pdf_utils, "_page_render_digest", recording_digest)
    outcome = compile_document_bytes(
        _pdf_outline_fixture_for_same_page(),
        CompilationProfileRef(
            "context-engine-pdf-outline-v1",
            PDF_TEXT_OUTLINE_V1,
        ),
    )

    assert type(outcome) is ParsedDocument
    assert render_calls == 1
    assert len(wrappers) == 1
    assert wrappers[0].closed is True


def test_pdf_without_outline_is_an_honest_typed_refusal() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    output = io.BytesIO()
    writer.write(output)

    outcome = compile_document_bytes(
        output.getvalue(),
        CompilationProfileRef(
            "context-engine-pdf-outline-v1",
            PDF_TEXT_OUTLINE_V1,
        ),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.OUTLINE_UNAVAILABLE


def test_page_render_digest_is_identical_for_same_pixels_in_distinct_pdfs() -> None:
    first = _pdf_outline_fixture()
    second = first + b"\n% byte-distinct container\n"
    outcomes = tuple(
        compile_document_bytes(
            source,
            CompilationProfileRef(
                "context-engine-pdf-outline-v1",
                PDF_TEXT_OUTLINE_V1,
            ),
        )
        for source in (first, second)
    )

    assert all(type(outcome) is ParsedDocument for outcome in outcomes)
    documents = tuple(
        outcome for outcome in outcomes if type(outcome) is ParsedDocument
    )
    assert documents[0].artifact_digest != documents[1].artifact_digest
    assert documents[0].units is not None and documents[1].units is not None
    first_locator = documents[0].units[0].locators[0]
    second_locator = documents[1].units[0].locators[0]
    assert type(first_locator) is PdfRegionLocator
    assert type(second_locator) is PdfRegionLocator
    assert first_locator.page_render_digest == second_locator.page_render_digest
    assert documents[0].content_hash == documents[1].content_hash
    assert documents[0].compilation_digest != documents[1].compilation_digest


@pytest.mark.parametrize(
    ("profile_ref", "fixture"),
    (
        (DOCX_CONFIG_V1, _docx_fixture),
        (PDF_TEXT_OUTLINE_V1, _pdf_outline_fixture),
    ),
)
def test_profile_digest_is_identical_across_two_fresh_processes(
    profile_ref: str,
    fixture: object,
) -> None:
    assert callable(fixture)
    source = fixture()
    canonical_documents: list[bytes] = []
    for hash_seed, thread_count in (("17", "1"), ("941", "2")):
        environment = _document_runner_environment(
            hash_seed=hash_seed,
            thread_count=thread_count,
        )
        completed = subprocess.run(
            [
                sys.executable,
                "-m",
                "applications.document_compiler_runner",
                "--profile",
                profile_ref,
            ],
            input=source,
            capture_output=True,
            check=True,
            cwd=REPOSITORY_ROOT,
            env=environment,
            timeout=30,
        )
        envelope = json.loads(completed.stdout)
        assert envelope["outcome"] == "parsed"
        canonical_documents.append(
            base64.b64decode(envelope["document"], validate=True)
        )

    assert canonical_documents[0] == canonical_documents[1]
    first = deserialize_parsed_document(canonical_documents[0])
    second = deserialize_parsed_document(canonical_documents[1])
    assert first.compilation_digest == second.compilation_digest


def test_local_runner_uses_the_shared_deterministic_environment(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    observed_environment: dict[str, str] | None = None

    def recording_run(*args: object, **kwargs: object) -> SimpleNamespace:
        del args
        nonlocal observed_environment
        observed_environment = cast(dict[str, str], kwargs["env"])
        return SimpleNamespace(
            returncode=0,
            stdout=json.dumps(
                {
                    "outcome": "failure",
                    "failure": {
                        "code": DocumentCompilationFailureCode.INVALID_ARTIFACT.value
                    },
                }
            ).encode("utf-8"),
        )

    monkeypatch.setattr(subprocess, "run", recording_run)
    outcome = compile_in_local_document_runner(
        BytesArtifactSource(b"invalid docx"),
        DOCX_CONFIG_V1,
        acceptance_context=acceptance_context(),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.INVALID_ARTIFACT
    assert observed_environment == _document_runner_environment()


@pytest.mark.parametrize("hash_seed", ("4294967296", "99999999999999999999"))
def test_document_runner_environment_rejects_out_of_range_hash_seed(
    hash_seed: str,
) -> None:
    with pytest.raises(ValueError, match="hash seed is out of range"):
        _document_runner_environment(hash_seed=hash_seed)


@pytest.mark.parametrize(
    ("field", "value"),
    (("hash_seed", "١"), ("thread_count", "１")),
)
def test_document_runner_environment_rejects_non_ascii_decimal_controls(
    field: str,
    value: str,
) -> None:
    with pytest.raises(ValueError, match="controls must be decimal integers"):
        _document_runner_environment(**{field: value})


@pytest.mark.parametrize("profile_ref", (DOCX_CONFIG_V1, PDF_TEXT_OUTLINE_V1))
def test_malformed_artifact_is_a_closed_typed_refusal(profile_ref: str) -> None:
    outcome = compile_in_local_document_runner(
        BytesArtifactSource(b"not the declared format"),
        profile_ref,
        acceptance_context=acceptance_context(),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.INVALID_ARTIFACT


def test_format_neutral_constructor_rejects_more_than_the_hard_unit_bound() -> None:
    digest = "0" * 64
    unit = StructuralUnit(
        ordinal=0,
        kind=DocumentStructuralKind.PARAGRAPH,
        text="bounded",
        locators=(
            DocxXmlLocator(
                artifact_digest=digest,
                part_uri="/word/document.xml",
                block_ordinal=0,
                xml_digest=digest,
            ),
        ),
    )
    units = tuple(
        StructuralUnit(
            ordinal=ordinal,
            kind=unit.kind,
            text=unit.text,
            locators=(
                DocxXmlLocator(
                    artifact_digest=digest,
                    part_uri="/word/document.xml",
                    block_ordinal=ordinal,
                    xml_digest=digest,
                ),
            ),
        )
        for ordinal in range(MAX_FORMAT_DOCUMENT_UNITS + 1)
    )

    with pytest.raises(ValueError, match="structural-unit hard bound"):
        ParsedDocument.format_neutral(
            artifact_digest=digest,
            profile=CompilationProfileRef(
                "context-engine-docx-v1",
                DOCX_CONFIG_V1,
            ),
            units=units,
        )


def test_format_text_bound_counts_copied_heading_ancestry(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import engine.supply.documents as document_contracts

    monkeypatch.setattr(document_contracts, "MAX_FORMAT_DOCUMENT_TEXT_CHARACTERS", 10)
    digest = "0" * 64
    units = (
        StructuralUnit(
            ordinal=0,
            kind=DocumentStructuralKind.HEADING,
            text="12345",
            locators=(DocxXmlLocator(digest, "/word/document.xml", 0, digest),),
            heading_level=1,
        ),
        StructuralUnit(
            ordinal=1,
            kind=DocumentStructuralKind.PARAGRAPH,
            text="1",
            locators=(DocxXmlLocator(digest, "/word/document.xml", 1, digest),),
            heading_ancestry=("12345",),
        ),
    )

    with pytest.raises(ValueError, match="text hard bound"):
        ParsedDocument.format_neutral(
            artifact_digest=digest,
            profile=CompilationProfileRef(
                "context-engine-docx-v1",
                DOCX_CONFIG_V1,
            ),
            units=units,
        )


def test_docx_compiler_types_ancestry_only_overflow_as_document_bound(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import engine.supply.documents as document_contracts

    monkeypatch.setattr(document_contracts, "MAX_FORMAT_DOCUMENT_TEXT_CHARACTERS", 10)
    monkeypatch.setattr(
        ragflow_document_adapter,
        "MAX_FORMAT_DOCUMENT_TEXT_CHARACTERS",
        10,
    )
    document = Document()
    document.add_heading("12345", level=1)
    document.add_paragraph("1")

    outcome = compile_document_bytes(
        _save_docx(document),
        CompilationProfileRef("context-engine-docx-v1", DOCX_CONFIG_V1),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.DOCUMENT_BOUND_EXCEEDED


def test_pdf_compiler_types_ancestry_only_overflow_as_document_bound(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import engine.supply.documents as document_contracts

    monkeypatch.setattr(document_contracts, "MAX_FORMAT_DOCUMENT_TEXT_CHARACTERS", 10)
    monkeypatch.setattr(
        ragflow_document_adapter,
        "MAX_FORMAT_DOCUMENT_TEXT_CHARACTERS",
        10,
    )
    monkeypatch.setattr(
        ragflow_document_adapter,
        "extract_pdf_outlines",
        lambda *_args, **_kwargs: (
            RawPdfOutline("12345", 0, 1, (0.0, 0.0, 1.0, 1.0), "0" * 64),
            RawPdfOutline("1", 1, 1, (0.0, 0.0, 1.0, 1.0), "0" * 64),
        ),
    )

    outcome = compile_document_bytes(
        _pdf_outline_fixture_for_same_page(),
        CompilationProfileRef(
            "context-engine-pdf-outline-v1",
            PDF_TEXT_OUTLINE_V1,
        ),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.DOCUMENT_BOUND_EXCEEDED


@pytest.mark.parametrize(
    "profile",
    (
        CompilationProfileRef("context-engine-pdf-outline-v1", DOCX_CONFIG_V1),
        CompilationProfileRef("unsupported-compiler-v1", DOCX_CONFIG_V1),
        CompilationProfileRef("context-engine-docx-v1", PDF_TEXT_OUTLINE_V1),
        CompilationProfileRef("unsupported-compiler-v1", PDF_TEXT_OUTLINE_V1),
    ),
)
def test_raw_compiler_refuses_unsupported_exact_profile_identity(
    profile: CompilationProfileRef,
) -> None:
    outcome = compile_document_bytes(_docx_fixture(), profile)

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.UNKNOWN_PROFILE


def test_format_constructor_rejects_wrong_compiler_and_cross_artifact_locator() -> (
    None
):
    artifact_digest = "1" * 64
    unit = StructuralUnit(
        ordinal=0,
        kind=DocumentStructuralKind.PARAGRAPH,
        text="bounded",
        locators=(
            DocxXmlLocator(
                artifact_digest="2" * 64,
                part_uri="/word/document.xml",
                block_ordinal=0,
                xml_digest="3" * 64,
            ),
        ),
    )
    with pytest.raises(ValueError, match="compiler/profile identity"):
        ParsedDocument.format_neutral(
            artifact_digest=artifact_digest,
            profile=CompilationProfileRef("forged-compiler", DOCX_CONFIG_V1),
            units=(unit,),
        )
    with pytest.raises(ValueError, match="bind the document artifact"):
        ParsedDocument.format_neutral(
            artifact_digest=artifact_digest,
            profile=CompilationProfileRef("context-engine-docx-v1", DOCX_CONFIG_V1),
            units=(unit,),
        )


def test_deserializer_rejects_forged_compiler_and_locator_artifact() -> None:
    compiled = compile_document_bytes(
        _docx_fixture(),
        CompilationProfileRef("context-engine-docx-v1", DOCX_CONFIG_V1),
    )
    assert type(compiled) is ParsedDocument
    canonical = json.loads(canonicalize_parsed_document(compiled))
    canonical["profile"]["compilerRef"] = "forged-compiler"
    with pytest.raises(ValueError, match="compiler/profile identity"):
        deserialize_parsed_document(rfc8785.dumps(canonical))
    canonical["profile"]["compilerRef"] = "context-engine-docx-v1"
    canonical["units"][0]["unexpected"] = "must refuse"
    with pytest.raises(ValueError, match="unit has unexpected fields"):
        deserialize_parsed_document(rfc8785.dumps(canonical))
    del canonical["units"][0]["unexpected"]
    canonical["units"][0]["locators"][0]["artifactDigest"] = "f" * 64
    with pytest.raises(ValueError, match="bind the document artifact"):
        deserialize_parsed_document(rfc8785.dumps(canonical))


def test_unleased_document_runner_requires_private_acceptance_capability() -> None:
    unchecked_runner = cast(Any, compile_in_local_document_runner)
    outcome = unchecked_runner(
        BytesArtifactSource(_docx_fixture()),
        DOCX_CONFIG_V1,
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.RUNNER_UNAVAILABLE


def test_raw_compiler_converts_constructor_rejection_to_typed_failure(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def reject_constructor(*args: object, **kwargs: object) -> ParsedDocument:
        raise ValueError("domain constructor rejected parser output")

    monkeypatch.setattr(ParsedDocument, "format_neutral", reject_constructor)
    outcome = compile_document_bytes(
        _docx_fixture(),
        CompilationProfileRef("context-engine-docx-v1", DOCX_CONFIG_V1),
    )

    assert type(outcome) is DocumentCompilationFailure
    assert outcome.code is DocumentCompilationFailureCode.INVALID_ARTIFACT
